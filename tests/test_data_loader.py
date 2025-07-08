"""
测试data_loader模块的功能
测试不同格式文档的加载和解析正确性
"""

from rag_core.data_loader import load_documents
import os


def test_load_txt(tmp_path):
    """
    测试TXT文件加载功能

    测试内容：
    1. 创建包含多段中文内容的TXT文件
    2. 验证文档加载结果
    3. 验证段落分割正确
    4. 验证内容完整性

    验证要点：
    - 返回类型为list
    - 段落数量正确
    - 每段内容完整
    """
    test_content = """第一段内容。

第二段内容。

第三段内容。"""
    test_file = tmp_path / "test.txt"
    test_file.write_text(test_content, encoding="utf-8")
    docs = load_documents(str(test_file))
    assert isinstance(docs, list)
    assert len(docs) == 3
    assert docs[0] == "第一段内容。"
    assert docs[1] == "第二段内容。"
    assert docs[2] == "第三段内容。"


def test_load_docx(tmp_path):
    """
    测试DOCX文件加载功能

    测试内容：
    1. 创建包含多段中文内容的DOCX文件
    2. 验证文档加载结果
    3. 验证段落分割正确
    4. 验证内容完整性

    验证要点：
    - 返回类型为list
    - 段落数量正确（跳过空段落）
    - 每段内容完整
    """
    from docx import Document

    test_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("第一段内容。")
    doc.add_paragraph("")  # 空段落
    doc.add_paragraph("第二段内容。")
    doc.add_paragraph("第三段内容。")
    doc.save(str(test_file))
    docs = load_documents(str(test_file))
    assert isinstance(docs, list)
    assert len(docs) == 3
    assert docs[0] == "第一段内容。"
    assert docs[1] == "第二段内容。"
    assert docs[2] == "第三段内容。"


def test_load_pdf(tmp_path):
    """
    测试PDF文件加载功能

    测试内容：
    1. 创建包含中文内容的PDF文件
    2. 处理不同系统的字体路径
    3. 验证文档加载结果
    4. 验证内容提取正确

    验证要点：
    - 返回类型为list
    - 包含所有段落内容
    - 支持中文PDF解析

    注意：
    - 自动检测系统字体路径
    - 支持macOS、Linux、Windows
    - 字体不可用时fallback到英文
    """
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase import pdfmetrics
    import sys

    test_file = tmp_path / "test.pdf"

    # 注册中文字体，支持多平台
    font_path = None
    for path in [
        "/System/Library/Fonts/STHeiti Medium.ttc",  # macOS
        "/usr/share/fonts/truetype/arphic/uming.ttc",  # Linux
        "C:/Windows/Fonts/simfang.ttf",  # Windows
    ]:
        if os.path.exists(path):
            font_path = path
            break

    if font_path:
        pdfmetrics.registerFont(TTFont("CustomChinese", font_path))
        font_name = "CustomChinese"
    else:
        # fallback: 用内置英文字体，测试英文
        font_name = "Helvetica"

    # 创建PDF文件
    c = canvas.Canvas(str(test_file), pagesize=A4)
    c.setFont(font_name, 14)
    c.drawString(100, 800, "第一段内容。")
    c.drawString(100, 780, "第二段内容。")
    c.drawString(100, 760, "第三段内容。")
    c.save()

    # 验证PDF加载结果
    docs = load_documents(str(test_file))
    # 只要包含所有段落即可（PDF解析可能有误差）
    assert isinstance(docs, list)
    assert any("第一段内容" in d for d in docs)
    assert any("第二段内容" in d for d in docs)
    assert any("第三段内容" in d for d in docs)
