"""
document_processor.py
增强版文档预处理器，支持多种文档格式和智能预处理功能。
"""

import os
import re
import json
import csv
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    增强版文档预处理器
    支持多种文档格式和智能预处理功能
    """

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        初始化文档预处理器

        :param config: 预处理配置
        """
        self.config = config or {}
        self.supported_formats = {
            ".txt": self._process_txt,
            ".md": self._process_markdown,
            ".docx": self._process_docx,
            ".pdf": self._process_pdf,
            ".html": self._process_html,
            ".json": self._process_json,
            ".csv": self._process_csv,
            ".xlsx": self._process_excel,
            ".xls": self._process_excel,
        }

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        处理单个文档

        :param file_path: 文档路径
        :return: 处理结果字典
        """
        try:
            file_path_obj = Path(file_path)
            if not file_path_obj.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")

            # 获取文件扩展名
            ext = file_path_obj.suffix.lower()

            # 检查是否支持该格式
            if ext not in self.supported_formats:
                raise ValueError(f"不支持的文件格式: {ext}")

            # 调用对应的处理函数
            processor = self.supported_formats[ext]
            result = processor(str(file_path))

            # 添加元数据
            result["metadata"] = {
                "file_path": str(file_path_obj),
                "file_name": file_path_obj.name,
                "file_size": file_path_obj.stat().st_size,
                "file_extension": ext,
                "processing_time": result.get("processing_time", 0),
            }

            return result

        except Exception as e:
            logger.error(f"处理文档失败 {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "content": "",
                "metadata": {
                    "file_path": str(file_path),
                    "file_name": (
                        Path(file_path).name if Path(file_path).exists() else "unknown"
                    ),
                },
            }

    def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """处理TXT文件"""
        import time

        start_time = time.time()

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # 应用文本清理
        content = self._clean_text(content)

        return {
            "success": True,
            "content": content,
            "processing_time": time.time() - start_time,
            "format": "txt",
        }

    def _process_markdown(self, file_path: str) -> Dict[str, Any]:
        """处理Markdown文件"""
        import time

        start_time = time.time()

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # 移除Markdown标记
        content = self._remove_markdown_syntax(content)

        # 应用文本清理
        content = self._clean_text(content)

        return {
            "success": True,
            "content": content,
            "processing_time": time.time() - start_time,
            "format": "markdown",
        }

    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """处理DOCX文件"""
        import time

        start_time = time.time()

        try:
            from docx import Document

            doc = Document(file_path)

            # 提取文本内容
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格内容
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # 合并内容
            content = "\n\n".join(paragraphs)

            # 添加表格内容
            if tables:
                table_text = []
                for i, table in enumerate(tables):
                    table_text.append(f"表格 {i+1}:")
                    for row in table:
                        table_text.append(" | ".join(row))
                    table_text.append("")
                content += "\n\n" + "\n".join(table_text)

            # 应用文本清理
            content = self._clean_text(content)

            return {
                "success": True,
                "content": content,
                "processing_time": time.time() - start_time,
                "format": "docx",
                "tables": tables,
            }

        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """处理PDF文件"""
        import time

        start_time = time.time()

        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)

            content_parts = []
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    # 添加页码标记
                    content_parts.append(f"[第{page_num + 1}页]\n{text}")

            content = "\n\n".join(content_parts)

            # 应用文本清理
            content = self._clean_text(content)

            return {
                "success": True,
                "content": content,
                "processing_time": time.time() - start_time,
                "format": "pdf",
                "page_count": len(reader.pages),
            }

        except ImportError:
            raise ImportError("请安装 PyPDF2: pip install PyPDF2")

    def _process_html(self, file_path: str) -> Dict[str, Any]:
        """处理HTML文件"""
        import time

        start_time = time.time()

        try:
            from bs4 import BeautifulSoup

            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()

            # 解析HTML
            soup = BeautifulSoup(html_content, "html.parser")

            # 移除脚本和样式标签
            for script in soup(["script", "style"]):
                script.decompose()

            # 提取文本
            text = soup.get_text()

            # 应用文本清理
            content = self._clean_text(text)

            return {
                "success": True,
                "content": content,
                "processing_time": time.time() - start_time,
                "format": "html",
            }

        except ImportError:
            raise ImportError("请安装 beautifulsoup4: pip install beautifulsoup4")

    def _process_json(self, file_path: str) -> Dict[str, Any]:
        """处理JSON文件"""
        import time

        start_time = time.time()

        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # 将JSON转换为文本
        content = self._json_to_text(data)

        # 应用文本清理
        content = self._clean_text(content)

        return {
            "success": True,
            "content": content,
            "processing_time": time.time() - start_time,
            "format": "json",
            "original_data": data,
        }

    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """处理CSV文件"""
        import time

        start_time = time.time()

        content_parts = []

        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)

            for row_num, row in enumerate(reader):
                if row_num == 0:  # 标题行
                    content_parts.append(f"列标题: {' | '.join(row)}")
                else:
                    content_parts.append(f"第{row_num}行: {' | '.join(row)}")

        content = "\n".join(content_parts)

        # 应用文本清理
        content = self._clean_text(content)

        return {
            "success": True,
            "content": content,
            "processing_time": time.time() - start_time,
            "format": "csv",
        }

    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """处理Excel文件"""
        import time

        start_time = time.time()

        try:
            import pandas as pd

            # 读取Excel文件
            excel_file = pd.ExcelFile(file_path)
            content_parts = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                content_parts.append(f"工作表: {sheet_name}")
                content_parts.append(f"列标题: {' | '.join(df.columns.astype(str))}")

                # 添加前几行数据作为示例
                for idx, row in df.head(10).iterrows():
                    row_data = " | ".join(row.astype(str))
                    content_parts.append(f"第{idx+1}行: {row_data}")

                content_parts.append("")

            content = "\n".join(content_parts)

            # 应用文本清理
            content = self._clean_text(content)

            return {
                "success": True,
                "content": content,
                "processing_time": time.time() - start_time,
                "format": "excel",
                "sheets": excel_file.sheet_names,
            }

        except ImportError:
            raise ImportError("请安装 pandas: pip install pandas")

    def _clean_text(self, text: str) -> str:
        """
        清理文本内容（保留段落分隔符\n\n，仅合并多余空格和多余空行）
        1. 合并多余空格（不影响换行）
        2. 合并多余空行为两个换行，保留段落分隔符
        3. 移除行首尾多余空格
        4. 移除页眉页脚、特殊字符、重复内容、标准化标点
        :param text: 原始文本
        :return: 清理后的文本
        """
        if not text:
            return ""
        # 合并多余空格（不影响换行）
        text = re.sub(r"[ \t]+", " ", text)
        # 合并多余空行为两个换行，保留段落分隔符
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 移除行首尾多余空格
        text = re.sub(r"^[ \t]+|[ \t]+$", "", text, flags=re.MULTILINE)
        # 移除页眉页脚（简单模式）
        text = self._remove_headers_footers(text)
        # 移除特殊字符
        text = self._remove_special_chars(text)
        # 标准化标点符号
        text = self._normalize_punctuation(text)
        # 移除重复内容（以段落为单位）
        text = self._remove_duplicates(text)
        return text.strip()

    def _remove_headers_footers(self, text: str) -> str:
        """移除页眉页脚"""
        # 移除页码
        text = re.sub(r"第\s*\d+\s*页", "", text)
        text = re.sub(r"Page\s*\d+", "", text)

        # 移除常见的页眉页脚内容
        header_footer_patterns = [
            r"机密文件|内部文件|保密文件",
            r"版权所有|Copyright|©",
            r"文档编号|Document ID|Doc ID",
            r"版本号|Version|V\d+\.\d+",
            r"创建日期|Created|Date:",
            r"最后修改|Last Modified|Updated:",
        ]

        for pattern in header_footer_patterns:
            text = re.sub(pattern, "", text, flags=re.IGNORECASE)

        return text

    def _remove_special_chars(self, text: str) -> str:
        """移除特殊字符"""
        # 保留中文字符、英文字母、数字、常用标点
        text = re.sub(
            r'[^\u4e00-\u9fff\w\s.,!?;:()（）【】""' "、。，！？；：]", "", text
        )
        return text

    def _normalize_punctuation(self, text: str) -> str:
        """标准化标点符号"""
        # 中英文标点转换
        punctuation_map = {
            "，": ",",
            "。": ".",
            "！": "!",
            "？": "?",
            "；": ";",
            "：": ":",
            "（": "(",
            "）": ")",
            '"': '"',
            '"': '"',
            """: "'", """: "'",
        }

        for chinese, english in punctuation_map.items():
            text = text.replace(chinese, english)

        return text

    def _remove_duplicates(self, text: str) -> str:
        """
        以段落为单位去重，保留分段结构
        :param text: 清洗后的文本
        :return: 去重后的文本
        """
        paragraphs = text.split("\n\n")
        seen = set()
        unique_paragraphs = []
        for para in paragraphs:
            para_stripped = para.strip()
            if para_stripped and para_stripped not in seen:
                seen.add(para_stripped)
                unique_paragraphs.append(para_stripped)
        return "\n\n".join(unique_paragraphs)

    def _remove_markdown_syntax(self, text: str) -> str:
        """移除Markdown语法"""
        # 移除标题标记
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)

        # 移除粗体和斜体标记
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\*(.*?)\*", r"\1", text)
        text = re.sub(r"__(.*?)__", r"\1", text)
        text = re.sub(r"_(.*?)_", r"\1", text)

        # 移除代码标记
        text = re.sub(r"`(.*?)`", r"\1", text)
        text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)

        # 移除链接
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

        # 移除列表标记
        text = re.sub(r"^[\s]*[-*+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^[\s]*\d+\.\s+", "", text, flags=re.MULTILINE)

        return text

    def _json_to_text(self, data: Any, indent: int = 0) -> str:
        """将JSON数据转换为文本"""
        if isinstance(data, dict):
            lines = []
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{'  ' * indent}{key}:")
                    lines.append(self._json_to_text(value, indent + 1))
                else:
                    lines.append(f"{'  ' * indent}{key}: {value}")
            return "\n".join(lines)
        elif isinstance(data, list):
            lines = []
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{'  ' * indent}项目 {i+1}:")
                    lines.append(self._json_to_text(item, indent + 1))
                else:
                    lines.append(f"{'  ' * indent}项目 {i+1}: {item}")
            return "\n".join(lines)
        else:
            return str(data)

    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式列表"""
        return list(self.supported_formats.keys())

    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """获取文档信息"""
        try:
            file_path_obj = Path(file_path)
            if not file_path_obj.exists():
                return {"error": "文件不存在"}

            ext = file_path_obj.suffix.lower()
            stats = file_path_obj.stat()

            info = {
                "file_name": file_path_obj.name,
                "file_path": str(file_path_obj),
                "file_size": stats.st_size,
                "file_extension": ext,
                "is_supported": ext in self.supported_formats,
                "last_modified": stats.st_mtime,
            }

            return info

        except Exception as e:
            return {"error": str(e)}


def process_documents(
    file_paths: List[str], config: Optional[Dict[str, Any]] = None
) -> List[Dict[str, Any]]:
    """
    批量处理文档

    :param file_paths: 文档路径列表
    :param config: 预处理配置
    :return: 处理结果列表
    """
    processor = DocumentProcessor(config)
    results = []

    for file_path in file_paths:
        result = processor.process_document(file_path)
        results.append(result)

    return results
